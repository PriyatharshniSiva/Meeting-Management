from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login as auth_login, logout as auth_logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import Meeting, ActionItem, CustomUser

def login_view(request):
    if request.method == 'GET' and request.user.is_authenticated:
        auth_logout(request)
        
    if request.method == 'POST':
        login_input = request.POST.get('email', '').strip() or request.POST.get('username', '').strip()
        p = request.POST.get('password', '').strip()
        
        if not login_input or not p:
            return render(request, 'login.html', {'error': 'Please provide both email/username and password.'})
            
        # Check if login_input is email
        username_to_auth = login_input
        if '@' in login_input:
            user_obj = CustomUser.objects.filter(email__iexact=login_input).first()
            if user_obj:
                username_to_auth = user_obj.username
                
        user = authenticate(request, username=username_to_auth, password=p)
        if user is not None:
            auth_login(request, user)
            request.session.set_expiry(0) # Strictly expire when browser closes
            return redirect('dashboard')
        else:
            return render(request, 'login.html', {'error': 'Invalid credentials. Please check your email and password.'})
            
    return render(request, 'login.html')


def signup_view(request):
    if request.method == 'GET' and request.user.is_authenticated:
        auth_logout(request)

        
    if request.method == 'POST':
        full_name = request.POST.get('full_name', '').strip()
        role = request.POST.get('role', 'EMPLOYEE').strip()
        email = request.POST.get('email', '').strip()
        phone_number = request.POST.get('phone_number', '').strip()
        password = request.POST.get('password', '').strip()
        confirm_password = request.POST.get('confirm_password', '').strip()
        
        if not full_name or not email or not password or not confirm_password:
            return render(request, 'signup.html', {'error': 'Full Name, Work Email, Password, and Confirm Password are required.'})
            
        if password != confirm_password:
            return render(request, 'signup.html', {'error': 'Passwords do not match. Please try again.'})
            
        if CustomUser.objects.filter(email__iexact=email).exists() or CustomUser.objects.filter(username__iexact=email).exists():
            return render(request, 'signup.html', {'error': 'An account with this email already exists.'})
            
        name_parts = full_name.split(' ', 1)
        first_name = name_parts[0]
        last_name = name_parts[1] if len(name_parts) > 1 else ''
        
        user = CustomUser.objects.create_user(
            username=email,
            email=email,
            password=password,
            first_name=first_name,
            last_name=last_name,
            phone_number=phone_number,
            role=role
        )
        messages.success(request, "Registration successful! Please log in with your credentials.")
        return redirect('login')
        
    return render(request, 'signup.html')


def logout_view(request):
    auth_logout(request)
    request.session.flush()
    return redirect('login')


@login_required
def dashboard_view(request):
    from django.utils import timezone
    import datetime
    
    now = timezone.now()
    
    all_meetings = Meeting.objects.all()
    total_meetings = all_meetings.count()
    upcoming_meetings = 0
    completed_meetings = 0
    
    for m in all_meetings:
        try:
            m_dt_naive = datetime.datetime.combine(m.date, m.time)
            m_dt = timezone.make_aware(m_dt_naive)
            if m_dt > now:
                upcoming_meetings += 1
            else:
                completed_meetings += 1
        except:
            pass
    
    pending_tasks = ActionItem.objects.filter(status='PENDING').count()
    overdue_tasks = 0
    for task in ActionItem.objects.filter(status='PENDING'):
        if task.due_date:
            try:
                task_dt = datetime.datetime.combine(task.due_date, datetime.time.min)
                task_dt = timezone.make_aware(task_dt)
                if task_dt < now:
                    overdue_tasks += 1
            except:
                pass
    
    from core.models import Participant, MeetingMinutes
    total_invites = Participant.objects.count()
    accepted_invites = Participant.objects.filter(rsvp_status='ACCEPTED').count()
    declined_invites = Participant.objects.filter(rsvp_status='DECLINED').count()
    attended_count = Participant.objects.filter(attendance_status='ATTENDED').count()
    
    total_transcriptions = MeetingMinutes.objects.exclude(history__isnull=True).exclude(history__exact='').count()
    ai_summaries = MeetingMinutes.objects.exclude(ai_summary__isnull=True).exclude(ai_summary__exact='').count()
    
    context = {
        'total_meetings': total_meetings,
        'upcoming_meetings': upcoming_meetings,
        'completed_meetings': completed_meetings,
        'pending_tasks': pending_tasks,
        'overdue_tasks': overdue_tasks,
        'total_invites': total_invites,
        'accepted_invites': accepted_invites,
        'declined_invites': declined_invites,
        'attended_count': attended_count,
        'total_transcriptions': total_transcriptions,
        'ai_summaries': ai_summaries,
    }
    return render(request, 'dashboard.html', context)

@login_required
def meetings_view(request):
    if request.method == 'POST':
        if request.user.role != 'ADMIN':
            from django.contrib import messages
            messages.error(request, "Only Admin users can create meetings.")
            return redirect('meetings')
            
        # Simple handler for creating a meeting
        title = request.POST.get('title')
        date = request.POST.get('date')
        time = request.POST.get('time')
        duration = request.POST.get('duration')
        m_type = request.POST.get('meeting_type')
        location = request.POST.get('location', '')
        meeting_link = request.POST.get('meeting_link', '')
        desc = request.POST.get('description')
        attachment = request.FILES.get('attachment')
        target_roles = request.POST.getlist('roles')
        
        if not date or not time:
            from django.contrib import messages
            messages.error(request, "Please provide a valid Date and Time format.")
            return redirect('meetings')
            
        # --- START ZOOM INTEGRATION ---
        if m_type == 'ONLINE':
            if not meeting_link:
                # Use the official Zoom Test Meeting so users don't get a 3001 Invalid ID error in the mock environment
                meeting_link = "https://zoom.us/test"
            
            desc += f"\\n\\n--- Meeting Link: {meeting_link} ---"
        # --- END ZOOM INTEGRATION ---
        
        meeting = Meeting.objects.create(
            title=title, date=date, time=time, 
            duration=duration, meeting_type=m_type, 
            location=location,
            meeting_link=meeting_link,
            description=desc, created_by=request.user,
            target_roles=target_roles,
            attachment=attachment
        )
        
        # Refresh to convert string date/time to python objects
        meeting.refresh_from_db()
        
        from core.models import CustomUser, Participant, InAppNotification
        
        target_user_ids = request.POST.getlist('users')
        external_emails_str = request.POST.get('external_emails', '')
        
        # If roles were used (fallback)
        if not target_user_ids and target_roles:
            target_user_ids = list(CustomUser.objects.filter(role__in=target_roles).values_list('id', flat=True))
            
        users_to_invite = CustomUser.objects.filter(id__in=target_user_ids)
        
        participants_to_create = []
        notifications = []
        
        # Internal Employees
        for u in users_to_invite:
            participants_to_create.append(Participant(meeting=meeting, user=u, rsvp_status='PENDING', attendance_status='PENDING'))
            notifications.append(InAppNotification(
                user=u, message=f"You have been invited to a new meeting: {title} on {date} at {time}",
                notification_type='INVITE', related_meeting=meeting
            ))
            
        # External Emails & Roles
        external_emails_list = request.POST.getlist('external_emails')
        external_roles_list = request.POST.getlist('external_roles')
        
        if external_roles_list:
            for role in external_roles_list:
                role_clean = role.strip()
                if role_clean:
                    participants_to_create.append(Participant(
                        meeting=meeting, 
                        external_email='', 
                        external_name=f"{role_clean} (External Guest)",
                        rsvp_status='PENDING', 
                        attendance_status='PENDING'
                    ))
        elif external_emails_list:
            for i, email in enumerate(external_emails_list):
                email_clean = email.strip()
                if email_clean:
                    role = external_roles_list[i].strip() if i < len(external_roles_list) else 'Guest'
                    name_prefix = email_clean.split('@')[0].replace('.', ' ').title()
                    display_name = f"{name_prefix} ({role})"
                    participants_to_create.append(Participant(
                        meeting=meeting, 
                        external_email=email_clean, 
                        external_name=display_name,
                        rsvp_status='PENDING', 
                        attendance_status='PENDING'
                    ))
        elif external_emails_str:
            import re
            emails = [email.strip() for email in re.split(r'[,\n]+', external_emails_str) if email.strip()]
            for email in emails:
                name_prefix = email.split('@')[0].replace('.', ' ').title()
                participants_to_create.append(Participant(
                    meeting=meeting, 
                    external_email=email, 
                    external_name=name_prefix,
                    rsvp_status='PENDING', 
                    attendance_status='PENDING'
                ))
                
        # Bulk Create
        created_participants = Participant.objects.bulk_create(participants_to_create)
        InAppNotification.objects.bulk_create(notifications)
        
        # We need the created participants with their IDs to pass to send_meeting_invites
        # bulk_create does not set PKs in SQLite in all versions, but let's query them
        all_participants = Participant.objects.filter(meeting=meeting)
        
        from core.utils import send_meeting_invites
        from django.contrib import messages
        import threading
        
        # Send emails in background
        def background_send(meeting_instance, participant_list):
            try:
                send_meeting_invites(meeting_instance, participant_list)
            except Exception as e:
                print(f"Error sending background emails: {e}")
                
        threading.Thread(target=background_send, args=(meeting, list(all_participants))).start()
        
        messages.success(request, f"Meeting created successfully! Invitations sent to {len(all_participants)} participants.")
        
        # Broadcast notification via WebSocket
        from asgiref.sync import async_to_sync
        from channels.layers import get_channel_layer
        channel_layer = get_channel_layer()
        async_to_sync(channel_layer.group_send)(
            'meeting_notifications',
            {
                'type': 'meeting_message',
                'message': f"New meeting created: {title} on {date} at {time}"
            }
        )
        
        # Check if meeting is happening right now (within next 5 minutes or already started today)
        from django.utils import timezone
        import datetime
        try:
            meeting_datetime = timezone.make_aware(datetime.datetime.strptime(f"{date} {time}", "%Y-%m-%d %H:%M"))
            now = timezone.now()
            # If the meeting is within 15 minutes of now (past or future), consider it "current time"
            if abs((meeting_datetime - now).total_seconds()) <= 15 * 60:
                if meeting_link:
                    return redirect(meeting_link)
        except Exception as e:
            pass
            
        return redirect('meetings')
        
    all_meetings = Meeting.objects.all().order_by('date', 'time')
    
    from django.utils import timezone
    import datetime
    now = timezone.now()
    
    meetings_list = []
    for m in all_meetings:
        try:
            m_dt_naive = datetime.datetime.combine(m.date, m.time)
            m_dt = timezone.make_aware(m_dt_naive)
            duration_delta = datetime.timedelta(minutes=m.duration if m.duration else 30)
            
            # Only append if the meeting hasn't ended yet
            if m_dt + duration_delta >= now:
                meetings_list.append(m)
        except Exception as e:
            print(f"Error calculating end_time for meeting {m.id}: {e}")
            # Fallback if datetime combine fails for any reason
            meetings_list.append(m)

    from core.models import CustomUser
    registered_users = CustomUser.objects.all().order_by('-date_joined')
    return render(request, 'meetings.html', {'meetings': meetings_list, 'users': registered_users})

@login_required
def invite_user_view(request, user_id):
    if request.user.role != 'ADMIN':
        from django.contrib import messages
        messages.error(request, "Only Admin users can invite users.")
        return redirect('meetings')
        
    from core.models import CustomUser
    from django.core.mail import send_mail
    from django.conf import settings
    from django.contrib import messages
    
    try:
        user_to_invite = CustomUser.objects.get(id=user_id)
        
        subject = "You've been invited to a Meeting!"
        message = f"Hello {user_to_invite.username},\n\nYou have been invited to join a meeting on MeetingMind. Please log in to view your meetings.\n\nBest,\nMeetingMind Team"
        
        send_mail(
            subject,
            message,
            settings.DEFAULT_FROM_EMAIL,
            [user_to_invite.email],
            fail_silently=False,
        )
        messages.success(request, f"Invite successfully sent to {user_to_invite.email}!")
    except CustomUser.DoesNotExist:
        messages.error(request, "User not found.")
    except Exception as e:
        print(f"EMAIL ERROR: {str(e)}")
        messages.error(request, f"Failed to send email to {user_to_invite.email}. Please check your SMTP settings in the .env file. Error: {str(e)}")
        
    return redirect('meetings')

@login_required
def delete_user_view(request, user_id):
    if request.user.role != 'ADMIN':
        from django.contrib import messages
        messages.error(request, "Only Admin users can delete users.")
        return redirect('meetings')
        
    from core.models import CustomUser
    from django.contrib import messages
    try:
        user_to_delete = CustomUser.objects.get(id=user_id)
        if user_to_delete != request.user:
            user_to_delete.delete()
            messages.success(request, f"User {user_to_delete.username} deleted successfully.")
        else:
            messages.error(request, "You cannot delete yourself.")
    except CustomUser.DoesNotExist:
        messages.error(request, "User not found.")
        
    return redirect('meetings')

@login_required
def delete_meeting_view(request, meeting_id):
    if request.user.role != 'ADMIN':
        from django.contrib import messages
        messages.error(request, "Only Admin users can delete meetings.")
        return redirect('meetings')
        
    from core.models import Meeting
    from django.contrib import messages
    try:
        meeting = Meeting.objects.get(id=meeting_id)
        meeting.delete()
        messages.success(request, f"Meeting '{meeting.title}' deleted successfully.")
    except Meeting.DoesNotExist:
        messages.error(request, "Meeting not found.")
        
    return redirect('meetings')

@login_required
def calendar_view(request):
    from core.models import Meeting
    from django.utils import timezone
    import datetime
    import json
    
    now = timezone.now()
    all_meetings = Meeting.objects.all().order_by('date', 'time')
    
    events = []
    upcoming = []
    for m in all_meetings:
        try:
            m_dt_naive = datetime.datetime.combine(m.date, m.time)
            m_dt = timezone.make_aware(m_dt_naive)
            duration = int(m.duration) if str(m.duration).isdigit() else 60
            end_dt = m_dt + datetime.timedelta(minutes=duration)
            
            # Use distinct colors based on meeting type or defaults
            class_name = 'event-primary'
            if m.meeting_type == 'ONLINE':
                class_name = 'event-client'
                
            events.append({
                'id': m.id,
                'title': m.title,
                'start': m_dt.isoformat(),
                'end': end_dt.isoformat(),
                'classNames': [class_name],
                'url': f"/mom/{m.id}/" if m.meeting_type == 'ONLINE' else ""
            })
            if m_dt > now:
                upcoming.append(m)
        except Exception as e:
            pass
            
    today_meetings = [m for m in all_meetings if m.date == now.date()]
    context = {
        'events_json': json.dumps(events),
        'upcoming_meetings': upcoming[:5],
        'total_meetings': len(all_meetings),
        'today_meetings': len(today_meetings),
        'upcoming_count': len(upcoming)
    }
    return render(request, 'calendar.html', context)

@login_required
def update_meeting_time(request):
    import json
    import datetime
    from django.http import JsonResponse
    from core.models import Meeting
    
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            meeting_id = data.get('meeting_id')
            new_datetime_str = data.get('new_datetime') # expected format: YYYY-MM-DDTHH:MM
            
            if not meeting_id or not new_datetime_str:
                return JsonResponse({'status': 'error', 'message': 'Invalid data'})
                
            meeting = Meeting.objects.get(id=meeting_id)
            # Only allow organizer or admin to reschedule
            if meeting.created_by != request.user and request.user.role != 'ADMIN':
                return JsonResponse({'status': 'error', 'message': 'Unauthorized'})
                
            new_dt = datetime.datetime.fromisoformat(new_datetime_str)
            meeting.date = new_dt.date()
            meeting.time = new_dt.time()
            meeting.save()
            
            return JsonResponse({'status': 'success'})
        except Meeting.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Meeting not found'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})
            
    return JsonResponse({'status': 'error', 'message': 'Invalid request'})

@login_required
def attendance_view(request):
    from core.models import Participant
    if request.user.role == 'ADMIN':
        participants = Participant.objects.all().order_by('-meeting__date', '-meeting__time')
    else:
        participants = Participant.objects.filter(user=request.user).order_by('-meeting__date', '-meeting__time')
        
    total = participants.count()
    present = participants.filter(attendance_status='ATTENDED').count()
    absent = participants.filter(attendance_status='ABSENT').count()
    late = participants.filter(attendance_status='LATE').count()
    
    return render(request, 'attendance.html', {
        'participants': participants,
        'total': total,
        'present': present,
        'absent': absent,
        'late': late,
    })

@login_required
def update_attendance(request, participant_id):
    if request.method == 'POST' and request.user.role == 'ADMIN':
        from core.models import Participant
        from django.contrib import messages
        try:
            p = Participant.objects.get(id=participant_id)
            new_status = request.POST.get('status')
            if new_status in dict(Participant._meta.get_field('attendance_status').choices):
                p.attendance_status = new_status
                p.save()
                messages.success(request, f"Updated attendance for {p.user.username}.")
        except Participant.DoesNotExist:
            messages.error(request, "Participant not found.")
    return redirect('attendance')

@login_required
def mom_view(request, meeting_id=None):
    from core.models import Meeting, MeetingMinutes
    import json
    
    if request.user.role == 'ADMIN':
        meetings = Meeting.objects.all().order_by('-date', '-time')
    else:
        meetings = Meeting.objects.filter(participants__user=request.user).order_by('-date', '-time')
        
    print(f"MOM VIEW Debug: User={request.user.username}, Role={request.user.role}, MeetingsCount={meetings.count()}")
        
    selected_meeting = None
    transcript_lines = []
    is_ongoing = False
    
    from django.utils import timezone
    import datetime
    
    if not meeting_id:
        now = timezone.localtime(timezone.now())
        print(f"MOM VIEW: NOW IS {now}")
        for m in meetings:
            start_dt = datetime.datetime.combine(m.date, m.time)
            if timezone.is_naive(start_dt):
                start_dt = timezone.make_aware(start_dt, timezone.get_default_timezone())
            try:
                duration = int(m.duration)
            except:
                duration = 60
            end_dt = start_dt + datetime.timedelta(minutes=duration)
            print(f"MOM VIEW Checking Meeting ID {m.id} | {m.title} | start: {start_dt} | end: {end_dt} | now: {now} | is_ongoing: {start_dt <= now <= end_dt}")
            if start_dt <= now <= end_dt:
                from django.shortcuts import redirect
                return redirect('mom_detail', meeting_id=m.id)
    
    if meeting_id:
        from django.shortcuts import get_object_or_404
        selected_meeting = get_object_or_404(Meeting, id=meeting_id)
        
        minutes, _ = MeetingMinutes.objects.get_or_create(meeting=selected_meeting)
                
        # Calculate if currently ongoing
        start_dt = datetime.datetime.combine(selected_meeting.date, selected_meeting.time)
        if timezone.is_naive(start_dt):
            start_dt = timezone.make_aware(start_dt, timezone.get_default_timezone())
        try:
            duration = int(selected_meeting.duration)
        except:
            duration = 60
        end_dt = start_dt + datetime.timedelta(minutes=duration)
        now = timezone.localtime(timezone.now())
        if start_dt <= now <= end_dt:
            is_ongoing = True
            
        # Try to load existing parsed transcript
        if minutes.history:
            try:
                transcript_lines = json.loads(minutes.history)
                speaker_colors = {}
                colors = ['#ef4444', '#3b82f6', '#22c55e', '#f59e0b', '#a855f7', '#ec4899', '#06b6d4', '#f97316', '#84cc16', '#6366f1']
                c_idx = 0
                for line in transcript_lines:
                    spk = line.get('speaker', 'Unknown')
                    if spk not in speaker_colors:
                        speaker_colors[spk] = colors[c_idx % len(colors)]
                        c_idx += 1
                    line['color'] = speaker_colors[spk]
            except:
                transcript_lines = []
        else:
            # If the meeting ended and no transcript was recorded
            if not is_ongoing and minutes.status == 'PENDING':
                minutes.status = 'NO_TRANSCRIPT'
                minutes.save()

    return render(request, 'mom.html', {
        'meetings': meetings,
        'selected_meeting': selected_meeting,
        'transcript_lines': transcript_lines,
        'is_ongoing': is_ongoing,
        'minutes': selected_meeting.minutes if selected_meeting and hasattr(selected_meeting, 'minutes') else None
    })

@login_required
def mom_content_partial(request, meeting_id):
    from core.models import Meeting, MeetingMinutes
    import json
    from django.shortcuts import get_object_or_404
    
    selected_meeting = get_object_or_404(Meeting, id=meeting_id)
    minutes, _ = MeetingMinutes.objects.get_or_create(meeting=selected_meeting)
    
    transcript_lines = []
    if minutes.history:
        try:
            transcript_lines = json.loads(minutes.history)
            speaker_colors = {}
            colors = ['#ef4444', '#3b82f6', '#22c55e', '#f59e0b', '#a855f7', '#ec4899', '#06b6d4', '#f97316', '#84cc16', '#6366f1']
            c_idx = 0
            for line in transcript_lines:
                spk = line.get('speaker', 'Unknown')
                if spk not in speaker_colors:
                    speaker_colors[spk] = colors[c_idx % len(colors)]
                    c_idx += 1
                line['color'] = speaker_colors[spk]
        except:
            pass
            
    return render(request, 'mom_content.html', {
        'selected_meeting': selected_meeting,
        'transcript_lines': transcript_lines,
        'minutes': minutes
    })

@login_required
def tasks_view(request):
    tasks = ActionItem.objects.all()
    return render(request, 'tasks.html', {'tasks': tasks})

@login_required
def reports_view(request):
    from core.models import Meeting
    from django.utils import timezone
    import datetime
    
    now = timezone.now()
    if request.user.role == 'ADMIN':
        all_meetings = Meeting.objects.all().order_by('-date', '-time')
    else:
        all_meetings = Meeting.objects.filter(participants__user=request.user).order_by('-date', '-time')
        
    completed_meetings = []
    
    for m in all_meetings:
        try:
            m_dt_naive = datetime.datetime.combine(m.date, m.time)
            m_dt = timezone.make_aware(m_dt_naive)
            duration = int(m.duration) if str(m.duration).isdigit() else 60
            end_dt = m_dt + datetime.timedelta(minutes=duration)
            if end_dt < now:
                completed_meetings.append(m)
        except Exception as e:
            pass
            
    return render(request, 'reports.html', {'completed_meetings': completed_meetings})

@login_required
def ai_features_view(request):
    return render(request, 'ai_features.html')

@login_required
def notifications_view(request):
    from core.models import InAppNotification
    notifications = request.user.notifications.all().order_by('-created_at')
    
    if request.method == 'POST':
        action = request.POST.get('action')
        if action == 'mark_all_read':
            notifications.update(is_read=True)
            return redirect('notifications')
            
    return render(request, 'notifications.html', {'notifications': notifications})

@login_required
def rsvp_view(request, meeting_id):
    if request.method == 'POST':
        from core.models import Participant, InAppNotification
        from django.contrib import messages
        from django.utils import timezone
        
        status = request.POST.get('status') # ACCEPTED, DECLINED, TENTATIVE
        try:
            participant = Participant.objects.get(meeting_id=meeting_id, user=request.user)
            if status in ['ACCEPTED', 'DECLINED', 'TENTATIVE']:
                participant.rsvp_status = status
                participant.rsvp_time = timezone.now()
                participant.save()
                messages.success(request, f"RSVP updated to {status}.")
                
                # Notify Organizer
                InAppNotification.objects.create(
                    user=participant.meeting.created_by,
                    message=f"{request.user.username} has {status.lower()} the invitation for '{participant.meeting.title}'.",
                    notification_type='SYSTEM',
                    related_meeting=participant.meeting
                )
                
                # Broadcast WebSocket Notification
                from asgiref.sync import async_to_sync
                from channels.layers import get_channel_layer
                channel_layer = get_channel_layer()
                async_to_sync(channel_layer.group_send)(
                    'meeting_notifications',
                    {
                        'type': 'meeting_message',
                        'message': f"{request.user.username} has {status.lower()} your meeting invitation."
                    }
                )
        except Participant.DoesNotExist:
            messages.error(request, "You are not a participant of this meeting.")
            
    return redirect(request.META.get('HTTP_REFERER', 'notifications'))

def rsvp_email_view(request, meeting_id):
    from core.models import Participant, InAppNotification
    from django.utils import timezone
    from django.http import HttpResponse
    from django.shortcuts import redirect
    import datetime
    
    participant_id = request.GET.get('participant_id')
    status = request.GET.get('status')
    
    if not participant_id or status not in ['ACCEPTED', 'DECLINED']:
        return HttpResponse("Invalid request parameters.", status=400)
        
    try:
        participant = Participant.objects.get(id=participant_id, meeting_id=meeting_id)
        
        # Prevent duplicate responses
        if participant.rsvp_status != 'PENDING':
            status_text = participant.get_rsvp_status_display()
            html = f"""
            <!DOCTYPE html>
            <html>
                <head>
                    <meta name="viewport" content="width=device-width, initial-scale=1">
                    <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
                    <script>
                        document.addEventListener("DOMContentLoaded", function() {{
                            Swal.fire({{
                                title: 'Already Responded',
                                text: 'You have already {status_text.lower()} this invitation.',
                                icon: 'info',
                                confirmButtonText: 'Close',
                                confirmButtonColor: '#4f46e5'
                            }}).then(() => {{
                                window.location.href = '/meetings/';
                            }});
                        }});
                    </script>
                </head>
                <body style='background-color: #f3f4f6;'></body>
            </html>
            """
            return HttpResponse(html)
            
        participant.rsvp_status = status
        participant.rsvp_time = timezone.now()
        participant.save()
        
        participant_name = participant.user.username if participant.user else participant.external_name
        
        # Notify Organizer
        InAppNotification.objects.create(
            user=participant.meeting.created_by,
            message=f"{participant_name} has {status.lower()} the invitation for '{participant.meeting.title}'.",
            notification_type='SYSTEM',
            related_meeting=participant.meeting
        )
        
        # Broadcast WebSocket
        from asgiref.sync import async_to_sync
        from channels.layers import get_channel_layer
        channel_layer = get_channel_layer()
        async_to_sync(channel_layer.group_send)(
            'meeting_notifications',
            {
                'type': 'meeting_message',
                'message': f"{participant_name} has {status.lower()} your meeting invitation."
            }
        )
        
        if status == 'ACCEPTED':
            m = participant.meeting
            m_dt_naive = datetime.datetime.combine(m.date, m.time)
            m_dt = timezone.make_aware(m_dt_naive)
            now = timezone.now()
            
            duration_delta = datetime.timedelta(minutes=m.duration if m.duration else 30)
            end_dt = m_dt + duration_delta
            
            # Auto-join if meeting is starting soon (within 15 mins) or currently ongoing
            if now >= (m_dt - datetime.timedelta(minutes=15)) and now <= end_dt:
                if m.meeting_link:
                    participant.join_time = timezone.now()
                    participant.attendance_status = 'ATTENDED'
                    participant.save()
                    return redirect(m.meeting_link)
            
            redirect_url = m.meeting_link if m.meeting_link else '/meetings/'
            html = f"""
            <!DOCTYPE html>
            <html>
                <head>
                    <meta name="viewport" content="width=device-width, initial-scale=1">
                    <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
                    <script>
                        document.addEventListener("DOMContentLoaded", function() {{
                            Swal.fire({{
                                title: 'Accepted!',
                                html: 'You have successfully accepted the invitation for <strong>{participant.meeting.title}</strong>.<br><br><b>Organizer:</b> {participant.meeting.created_by.username}<br><b>Time:</b> {participant.meeting.date} at {participant.meeting.time}',
                                icon: 'success',
                                confirmButtonText: 'Join / Continue',
                                confirmButtonColor: '#10b981',
                                timer: 4000,
                                timerProgressBar: true
                            }}).then(() => {{
                                window.location.href = "{redirect_url}";
                            }});
                        }});
                    </script>
                </head>
                <body style='background-color: #f3f4f6;'></body>
            </html>
            """
        else:
            html = f"""
            <!DOCTYPE html>
            <html>
                <head>
                    <meta name="viewport" content="width=device-width, initial-scale=1">
                    <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
                    <script>
                        document.addEventListener("DOMContentLoaded", function() {{
                            Swal.fire({{
                                title: 'Declined',
                                text: 'You have declined the invitation for {participant.meeting.title}.',
                                icon: 'error',
                                confirmButtonText: 'Close',
                                confirmButtonColor: '#ef4444'
                            }}).then(() => {{
                                window.close();
                                window.location.href = '/meetings/';
                            }});
                        }});
                    </script>
                </head>
                <body style='background-color: #f3f4f6;'></body>
            </html>
            """
            
        return HttpResponse(html)
        
    except Participant.DoesNotExist:
        return HttpResponse("Participant or meeting not found.", status=404)

@login_required
def join_meeting(request, meeting_id):
    from core.models import Participant, Meeting, InAppNotification
    from django.utils import timezone
    
    try:
        meeting = Meeting.objects.get(id=meeting_id)
        participant = Participant.objects.get(meeting=meeting, user=request.user)
        
        participant.join_time = timezone.now()
        participant.attendance_status = 'ATTENDED'
        participant.save()
        
        # Notify Organizer
        InAppNotification.objects.create(
            user=meeting.created_by,
            message=f"{request.user.username} has joined the meeting '{meeting.title}'.",
            notification_type='SYSTEM',
            related_meeting=meeting
        )
        
        if meeting.meeting_link:
            return redirect(meeting.meeting_link)
    except Exception as e:
        pass
        
    return redirect('meetings')

@login_required
def leave_meeting(request, meeting_id):
    from core.models import Participant, Meeting, InAppNotification
    from django.utils import timezone
    
    try:
        meeting = Meeting.objects.get(id=meeting_id)
        participant = Participant.objects.get(meeting=meeting, user=request.user)
        
        participant.leave_time = timezone.now()
        participant.save()
        
        # Notify Organizer
        InAppNotification.objects.create(
            user=meeting.created_by,
            message=f"{request.user.username} has left the meeting '{meeting.title}'.",
            notification_type='SYSTEM',
            related_meeting=meeting
        )
    except Exception as e:
        pass
        
    return redirect('meetings')

@login_required
def settings_view(request):
    from core.models import CustomUser
    from django.contrib import messages
    if request.method == 'POST':
        action = request.POST.get('action')
        
        if action == 'update_profile':
            username = request.POST.get('username')
            email = request.POST.get('email')
            
            if username and email:
                if CustomUser.objects.filter(username=username).exclude(id=request.user.id).exists():
                    messages.error(request, "Username is already taken.")
                elif CustomUser.objects.filter(email=email).exclude(id=request.user.id).exists():
                    messages.error(request, "Email is already in use.")
                else:
                    request.user.username = username
                    request.user.email = email
                    request.user.save()
                    messages.success(request, "Profile updated successfully.")
            else:
                messages.error(request, "Username and Email cannot be empty.")
                
        elif action == 'update_password':
            current_password = request.POST.get('current_password')
            new_password = request.POST.get('new_password')
            confirm_password = request.POST.get('confirm_password')
            
            if request.user.check_password(current_password):
                if new_password == confirm_password:
                    if len(new_password) >= 6:
                        request.user.set_password(new_password)
                        request.user.save()
                        from django.contrib.auth import update_session_auth_hash
                        update_session_auth_hash(request, request.user)
                        messages.success(request, "Password updated successfully.")
                    else:
                        messages.error(request, "New password must be at least 6 characters.")
                else:
                    messages.error(request, "New passwords do not match.")
            else:
                messages.error(request, "Current password is incorrect.")
                
        elif action == 'update_smtp':
            if request.user.role != 'ADMIN':
                messages.error(request, "Only administrators can update SMTP settings.")
            else:
                import os
                from pathlib import Path
                env_path = Path(settings.BASE_DIR) / '.env'
                
                # Parse existing .env lines
                env_dict = {}
                if env_path.exists():
                    with open(env_path, 'r') as f:
                        for line in f:
                            if '=' in line and not line.strip().startswith('#'):
                                k, v = line.strip().split('=', 1)
                                env_dict[k] = v
                
                # Update dict with new values
                env_dict['EMAIL_HOST'] = request.POST.get('email_host', 'smtp.gmail.com')
                env_dict['EMAIL_PORT'] = request.POST.get('email_port', '587')
                env_dict['EMAIL_USE_TLS'] = 'True' if request.POST.get('email_use_tls') == 'on' else 'False'
                env_dict['EMAIL_HOST_USER'] = request.POST.get('email_host_user', '')
                if request.POST.get('email_host_password'):
                    env_dict['EMAIL_HOST_PASSWORD'] = request.POST.get('email_host_password')
                
                # Write back to .env
                with open(env_path, 'w') as f:
                    for k, v in env_dict.items():
                        f.write(f"{k}={v}\n")
                        
                messages.success(request, "SMTP settings updated in .env. You may need to restart the server for changes to take effect.")
                
    import os
    smtp_context = {
        'EMAIL_HOST': os.getenv('EMAIL_HOST', 'smtp.gmail.com'),
        'EMAIL_PORT': os.getenv('EMAIL_PORT', '587'),
        'EMAIL_USE_TLS': str(os.getenv('EMAIL_USE_TLS', 'True')).lower() in ['true', '1'],
        'EMAIL_HOST_USER': os.getenv('EMAIL_HOST_USER', ''),
    }
    return render(request, 'settings.html', {'smtp': smtp_context})

@login_required
def projects_view(request):
    from core.models import Project, ProjectAssignmentLog, CustomUser, InAppNotification
    from django.db.models import Count, Q
    from django.contrib import messages
    from django.core.mail import send_mail
    from django.conf import settings
    from django.urls import reverse
    import datetime

    if request.method == 'POST' and request.user.role == 'ADMIN':
        name = request.POST.get('name')
        description = request.POST.get('description')
        skills = request.POST.get('skills')
        team_size = request.POST.get('team_size')
        deadline = request.POST.get('deadline')
        priority = request.POST.get('priority', 'MEDIUM')
        assigned_user_ids = request.POST.getlist('assigned_users')
        
        try:
            team_size = int(team_size)
            project = Project.objects.create(
                name=name,
                description=description,
                required_skills=skills,
                required_team_size=team_size,
                deadline=deadline,
                priority=priority,
                created_by=request.user
            )
            
            # Find eligible employees (TL and EMPLOYEE)
            eligible_users = CustomUser.objects.filter(id__in=assigned_user_ids, role__in=['EMPLOYEE', 'TL'])
            
            logs = []
            notifications = []
            for u in eligible_users:
                logs.append(ProjectAssignmentLog(project=project, user=u, status='INVITED'))
                notifications.append(InAppNotification(
                    user=u,
                    message=f"New Project Invitation: {project.name}. Requires {project.required_team_size} members. Accept quickly!",
                    notification_type='PROJECT_INVITE',
                    related_project=project
                ))
            
            ProjectAssignmentLog.objects.bulk_create(logs)
            InAppNotification.objects.bulk_create(notifications)
            
            # Send Email Invitations
            subject = f"Invitation: {project.name}"
            for u in eligible_users:
                accept_url = request.build_absolute_uri(reverse('accept_project', args=[project.id]))
                decline_url = request.build_absolute_uri(reverse('decline_project', args=[project.id]))
                
                message = f"""Hello {u.username},
                
You have been invited to join a new project!

Project Name: {project.name}
Description: {project.description}
Required Skills: {project.required_skills}
Priority: {project.get_priority_display()}
Deadline: {project.deadline}

Please review the project details and accept or decline the invitation.

Accept Project: {accept_url}
Decline Project: {decline_url}
"""
                send_mail(subject, message, settings.DEFAULT_FROM_EMAIL, [u.email], fail_silently=True)
            
            messages.success(request, "Project created and invitations sent!")
        except Exception as e:
            messages.error(request, f"Error creating project: {str(e)}")
            
        return redirect('projects')
        
    context = {}
    if request.user.role == 'ADMIN':
        # Admin sees all their created projects with stats
        projects = Project.objects.filter(created_by=request.user).prefetch_related('assignment_logs__user').annotate(
            total_invites=Count('assignment_logs'),
            accepted_count=Count('assignment_logs', filter=Q(assignment_logs__status='ACCEPTED')),
            declined_count=Count('assignment_logs', filter=Q(assignment_logs__status='DECLINED'))
        ).order_by('-created_at')
        context['admin_projects'] = projects
        context['eligible_users'] = CustomUser.objects.filter(role__in=['EMPLOYEE', 'TL']).order_by('username')
    else:
        # Employees see projects they are invited to or accepted
        logs = ProjectAssignmentLog.objects.filter(user=request.user).select_related('project').order_by('-project__created_at')
        context['user_logs'] = logs

    return render(request, 'projects.html', context)

@login_required
def edit_project(request, project_id):
    from core.models import Project
    if request.method == 'POST' and request.user.role == 'ADMIN':
        try:
            project = Project.objects.get(id=project_id, created_by=request.user)
            project.name = request.POST.get('name')
            project.description = request.POST.get('description')
            project.required_skills = request.POST.get('skills')
            project.required_team_size = int(request.POST.get('team_size'))
            project.deadline = request.POST.get('deadline')
            project.priority = request.POST.get('priority', 'MEDIUM')
            project.save()
            messages.success(request, "Project updated successfully.")
        except Project.DoesNotExist:
            messages.error(request, "Project not found or permission denied.")
        except Exception as e:
            messages.error(request, f"Error updating project: {str(e)}")
    return redirect('projects')

@login_required
def delete_project(request, project_id):
    from core.models import Project
    if request.method == 'POST' and request.user.role == 'ADMIN':
        try:
            project = Project.objects.get(id=project_id, created_by=request.user)
            project.delete()
            messages.success(request, "Project deleted successfully.")
        except Project.DoesNotExist:
            messages.error(request, "Project not found or permission denied.")
    return redirect('projects')

@login_required
def accept_project(request, project_id):
    from core.models import Project, ProjectAssignmentLog, InAppNotification
    from django.db import transaction
    from django.contrib import messages
    from django.utils import timezone
    
    with transaction.atomic():
        try:
            # Lock the project row for concurrency
            project = Project.objects.select_for_update().get(id=project_id)
            
            if project.status == 'TEAM_FINALIZED':
                messages.error(request, "Project Team is already Full.")
                return redirect('projects')
                
            log = ProjectAssignmentLog.objects.get(project=project, user=request.user)
            
            if log.status != 'INVITED':
                messages.warning(request, f"You have already {log.status.lower()} this project.")
                return redirect('projects')
                
            # Count currently accepted members
            accepted_count = ProjectAssignmentLog.objects.filter(project=project, status='ACCEPTED').count()
            
            if accepted_count < project.required_team_size:
                log.status = 'ACCEPTED'
                log.timestamp = timezone.now()
                log.save()
                
                messages.success(request, f"You have successfully joined the project: {project.name}!")
                
                # Notify admin immediately about this specific acceptance
                InAppNotification.objects.create(
                    user=project.created_by,
                    message=f"{request.user.username} has ACCEPTED the project: {project.name}.",
                    notification_type='PROJECT_UPDATE',
                    related_project=project
                )

                # Check if we just filled the team
                if accepted_count + 1 >= project.required_team_size:
                    project.status = 'TEAM_FINALIZED'
                    project.save()
                    
                    # Notify admin
                    accepted_logs = ProjectAssignmentLog.objects.filter(project=project, status='ACCEPTED').order_by('timestamp')
                    names_and_times = "\\n".join([f"- {l.user.username} at {l.timestamp.strftime('%Y-%m-%d %H:%M:%S')}" for l in accepted_logs])
                    InAppNotification.objects.create(
                        user=project.created_by,
                        message=f"Team Finalized for project '{project.name}' (Size: {project.required_team_size}).\\nMembers:\\n{names_and_times}",
                        notification_type='PROJECT_UPDATE',
                        related_project=project
                    )
                    
                    # Notify remaining users that the project is no longer available
                    remaining_logs = ProjectAssignmentLog.objects.filter(project=project, status='INVITED')
                    full_notifications = []
                    for r_log in remaining_logs:
                        full_notifications.append(InAppNotification(
                            user=r_log.user,
                            message=f"The project '{project.name}' is no longer available. The team is full.",
                            notification_type='SYSTEM',
                            related_project=project
                        ))
                    InAppNotification.objects.bulk_create(full_notifications)
            else:
                # Due to a race condition that wasn't caught by the first check, it's actually full
                messages.error(request, "Project Team is already Full.")
                
        except Project.DoesNotExist:
            messages.error(request, "Project not found.")
        except ProjectAssignmentLog.DoesNotExist:
            messages.error(request, "You are not invited to this project.")
            
    return redirect('projects')

@login_required
def decline_project(request, project_id):
    from core.models import ProjectAssignmentLog, InAppNotification
    from django.contrib import messages
    from django.utils import timezone
    
    try:
        log = ProjectAssignmentLog.objects.get(project_id=project_id, user=request.user)
        if log.status == 'INVITED':
            log.status = 'DECLINED'
            log.timestamp = timezone.now()
            log.save()
            messages.success(request, "You have declined the project.")
            
            InAppNotification.objects.create(
                user=log.project.created_by,
                message=f"{request.user.username} has DECLINED the project: {log.project.name}.",
                notification_type='PROJECT_UPDATE',
                related_project=log.project
            )
        else:
            messages.warning(request, "You cannot decline at this stage.")
    except ProjectAssignmentLog.DoesNotExist:
        messages.error(request, "You are not invited to this project.")
        
    return redirect('projects')

def request_otp_view(request):
    if request.method == 'POST':
        email = request.POST.get('email', '').strip()
        from core.models import CustomUser, PasswordResetOTP
        from django.core.mail import send_mail
        from django.conf import settings
        import random
        import string
        
        user = CustomUser.objects.filter(email__iexact=email).first()
        if user:
            otp = ''.join(random.choices(string.digits, k=6))
            PasswordResetOTP.objects.create(user=user, otp=otp)
            
            subject = "Password Reset OTP - MeetingMind"
            message = f"Hello {user.get_full_name() or user.username},\n\nYour 6-digit OTP for resetting your MeetingMind password is: {otp}\n\nThis code will expire in 15 minutes.\n\nIf you did not request this, please ignore this email."
            
            try:
                send_mail(subject, message, settings.DEFAULT_FROM_EMAIL, [email], fail_silently=False)
            except Exception as e:
                import logging
                logging.error(f"Failed to send email to {email}: {str(e)}")

            
            request.session['reset_email'] = email
            messages.success(request, f"Verification OTP sent to {email}. Please check your email inbox.")
            return redirect('verify_otp')



        else:
            return render(request, 'registration/password_reset_form.html', {'error': 'No account found with this email address.'})
            
    return render(request, 'registration/password_reset_form.html')


def verify_otp_view(request):
    email = request.session.get('reset_email')
    if not email:
        return redirect('password_reset')
        
    if request.method == 'POST':
        otp = request.POST.get('otp', '').strip()
        from core.models import CustomUser, PasswordResetOTP
        from django.utils import timezone
        import datetime
        
        user = CustomUser.objects.filter(email__iexact=email).first()
        if user:
            otp_record = PasswordResetOTP.objects.filter(user=user, otp=otp, is_used=False).order_by('-created_at').first()
            if otp_record:
                if timezone.now() - otp_record.created_at < datetime.timedelta(minutes=15):
                    otp_record.is_used = True
                    otp_record.save()
                    request.session['otp_verified'] = True
                    return redirect('set_new_password')
                else:
                    return render(request, 'registration/password_reset_done.html', {'error': 'OTP has expired. Please request a new one.'})
            else:
                return render(request, 'registration/password_reset_done.html', {'error': 'Invalid OTP code. Please check your email and try again.'})
        else:
            return render(request, 'registration/password_reset_done.html', {'error': 'Invalid session or user not found.'})
            
    return render(request, 'registration/password_reset_done.html')

def set_new_password_view(request):
    email = request.session.get('reset_email')
    verified = request.session.get('otp_verified')
    if not email or not verified:
        return redirect('password_reset')
        
    if request.method == 'POST':
        password = request.POST.get('password', '').strip()
        confirm_password = request.POST.get('confirm_password', '').strip()
        
        if password != confirm_password:
            return render(request, 'registration/password_reset_confirm.html', {'error': 'Passwords do not match.'})
            
        from core.models import CustomUser
        user = CustomUser.objects.filter(email__iexact=email).first()
        if user:
            user.set_password(password)
            user.save()
            
            if 'reset_email' in request.session: del request.session['reset_email']
            if 'otp_verified' in request.session: del request.session['otp_verified']
            
            messages.success(request, "Password reset complete! Please log in with your new password.")
            return redirect('login')
        else:
            return redirect('password_reset')
            
    return render(request, 'registration/password_reset_confirm.html')



@login_required
def generate_summary(request, meeting_id):
    from core.models import Meeting, MeetingMinutes, ActionItem, CustomUser
    from django.shortcuts import get_object_or_404, redirect
    from django.contrib import messages
    import json
    import os
    import google.generativeai as genai
    
    meeting = get_object_or_404(Meeting, id=meeting_id)
    if not hasattr(meeting, 'minutes') or not meeting.minutes.history:
        return JsonResponse({'status': 'error', 'message': "No transcript available to summarize."})
        
    minutes = meeting.minutes
    minutes.status = 'GENERATING'
    minutes.save()
    
    try:
        transcript_lines = json.loads(minutes.history)
    except:
        transcript_lines = []
        
    transcript_text = "\n".join([f"{t.get('speaker', 'Unknown')}: {t.get('text', '')}" for t in transcript_lines])
    
    # Check if we have anything to summarize
    if len(transcript_text.strip()) < 10:
        minutes.status = 'NO_TRANSCRIPT'
        minutes.save()
        return JsonResponse({'status': 'error', 'message': "Transcript is too short."})
        
    api_key = os.getenv("GEMINI_API_KEY", "")
    if not api_key:
        minutes.status = 'PENDING'
        minutes.save()
        return JsonResponse({'status': 'error', 'message': "GEMINI_API_KEY is not configured in .env"})
        
    genai.configure(api_key=api_key)
    
    prompt = f"""
    You are an expert executive assistant. Review the following meeting transcript.
    Extract the following structured information and return it EXACTLY as a valid JSON object without markdown formatting blocks (do not wrap in ```json).
    {{
        "Summary": "A concise paragraph summarizing the overall meeting.",
        "Discussion_Points": ["Point 1", "Point 2"],
        "Decisions_Taken": ["Decision 1", "Decision 2"],
        "Action_Items": [
            {{"task": "Task description", "owner": "Assigned Person", "due_date": "YYYY-MM-DD or None"}}
        ]
    }}
    
    Transcript:
    {transcript_text}
    """
    
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt)
        result_text = response.text.strip()
        
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
            
        parsed = json.loads(result_text)
        
        # Format the discussion points into the AI summary text
        summary = parsed.get("Summary", "No summary provided.")
        discussion_pts = parsed.get("Discussion_Points", [])
        if discussion_pts:
            summary += "\n\n### Discussion Points\n- " + "\n- ".join(discussion_pts)
            
        minutes.ai_summary = summary
        minutes.decisions = parsed.get("Decisions_Taken", [])
        minutes.action_items_raw = parsed.get("Action_Items", [])
        minutes.save()
        
        # Optional: Auto-create Action Items in DB
        for item in minutes.action_items_raw:
            owner_str = item.get("owner", "")
            # Try to map owner to existing user
            owner_user = CustomUser.objects.filter(username__icontains=owner_str).first()
            if owner_user:
                ActionItem.objects.create(
                    meeting=meeting,
                    task_name=item.get("task", "New Task")[:255],
                    description=item.get("task", ""),
                    assigned_to=owner_user,
                    due_date=meeting.date # Fallback to meeting date
                )
        
        # Now automatically generate exports and save to FileFields
        from io import BytesIO
        from django.core.files.base import ContentFile
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        import docx
        
        # PDF Generation
        try:
            pdf_buffer = BytesIO()
            p = canvas.Canvas(pdf_buffer, pagesize=letter)
            p.drawString(100, 750, f"Transcript: {meeting.title}")
            y = 730
            for line in transcript_lines:
                if y < 50:
                    p.showPage()
                    y = 750
                text = f"{line.get('speaker')}: {line.get('text')}"
                p.drawString(100, y, text[:90])
                y -= 15
            p.showPage()
            p.save()
            pdf_buffer.seek(0)
            minutes.pdf_export.save(f'transcript_{meeting.id}.pdf', ContentFile(pdf_buffer.read()), save=False)
        except Exception as e:
            print("Error generating PDF:", e)
            
        # DOCX Generation
        try:
            doc = docx.Document()
            doc.add_heading(f"Transcript: {meeting.title}", 0)
            for line in transcript_lines:
                p_doc = doc.add_paragraph()
                p_doc.add_run(f"{line.get('speaker')}: ").bold = True
                p_doc.add_run(f"{line.get('text')}")
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            minutes.docx_export.save(f'transcript_{meeting.id}.docx', ContentFile(docx_buffer.read()), save=False)
        except Exception as e:
            print("Error generating DOCX:", e)
            
        import datetime
        from django.utils import timezone
        now_str = timezone.now().strftime("%Y-%m-%dT%H:%M:%S")
        current_logs = minutes.processing_logs if isinstance(minutes.processing_logs, list) else []
        current_logs.append({'timestamp': now_str, 'event': 'Summary Generated', 'type': 'success'})
        current_logs.append({'timestamp': now_str, 'event': 'Minutes Created', 'type': 'success'})
        minutes.processing_logs = current_logs
            
        minutes.status = 'READY'
        minutes.save()
        
        return JsonResponse({'status': 'success'})
    except Exception as e:
        import datetime
        from django.utils import timezone
        now_str = timezone.now().strftime("%Y-%m-%dT%H:%M:%S")
        current_logs = minutes.processing_logs if isinstance(minutes.processing_logs, list) else []
        current_logs.append({'timestamp': now_str, 'event': f'Error: {str(e)}', 'type': 'error'})
        minutes.processing_logs = current_logs
        
        minutes.status = 'PENDING'
        minutes.save()
        return JsonResponse({'status': 'error', 'message': str(e)})

@login_required
def export_pdf(request, meeting_id):
    from django.http import HttpResponse
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from core.models import Meeting
    import json
    
    meeting = Meeting.objects.get(id=meeting_id)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="transcript_{meeting.id}.pdf"'
    
    p = canvas.Canvas(response, pagesize=letter)
    p.drawString(100, 750, f"Transcript: {meeting.title}")
    
    y = 730
    if hasattr(meeting, 'minutes') and meeting.minutes.history:
        try:
            lines = json.loads(meeting.minutes.history)
            for line in lines:
                if y < 50:
                    p.showPage()
                    y = 750
                # reportlab doesn't wrap text automatically with drawString, but this is a basic export
                text = f"{line.get('speaker')}: {line.get('text')}"
                p.drawString(100, y, text[:90]) # limit length to prevent runoff
                y -= 15
        except:
            pass
            
    p.showPage()
    p.save()
    return response

@login_required
def download_transcript(request, meeting_id):
    from django.http import HttpResponse
    from core.models import Meeting
    import json
    
    meeting = Meeting.objects.get(id=meeting_id)
    response = HttpResponse(content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename="transcript_{meeting.id}.txt"'
    
    if hasattr(meeting, 'minutes') and meeting.minutes.history:
        try:
            lines = json.loads(meeting.minutes.history)
            for line in lines:
                response.write(f"{line.get('speaker', 'Unknown')}: {line.get('text', '')}\n")
        except Exception as e:
            response.write(f"Error parsing transcript: {e}")
    else:
        response.write("No transcript available.")
        
    return response

@login_required
def download_ai_summary(request, meeting_id):
    from django.http import HttpResponse
    from core.models import Meeting
    import json
    
    meeting = Meeting.objects.get(id=meeting_id)
    response = HttpResponse(content_type='application/json')
    response['Content-Disposition'] = f'attachment; filename="ai_summary_{meeting.id}.json"'
    
    if hasattr(meeting, 'minutes'):
        data = {
            "title": meeting.title,
            "date": str(meeting.date),
            "summary": meeting.minutes.ai_summary,
            "decisions": meeting.minutes.decisions,
            "action_items": meeting.minutes.action_items_raw
        }
        response.write(json.dumps(data, indent=4))
    else:
        response.write(json.dumps({"error": "No minutes available."}))
        
    return response

@login_required
def export_docx(request, meeting_id):
    from django.http import HttpResponse
    import docx
    from core.models import Meeting
    import json
    import io
    
    meeting = Meeting.objects.get(id=meeting_id)
    doc = docx.Document()
    doc.add_heading(f"Transcript: {meeting.title}", 0)
    
    if hasattr(meeting, 'minutes') and meeting.minutes.history:
        try:
            lines = json.loads(meeting.minutes.history)
            for line in lines:
                p = doc.add_paragraph()
                p.add_run(f"{line.get('speaker')}: ").bold = True
                p.add_run(f"{line.get('text')}")
        except:
            pass
            
    f = io.BytesIO()
    doc.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="transcript_{meeting.id}.docx"'
    response['Content-Length'] = length
    return response

@login_required
def export_txt(request, meeting_id):
    from django.http import HttpResponse
    from core.models import Meeting
    import json
    
    meeting = Meeting.objects.get(id=meeting_id)
    content = f"Transcript: {meeting.title}\n\n"
    
    if hasattr(meeting, 'minutes') and meeting.minutes.history:
        try:
            lines = json.loads(meeting.minutes.history)
            for line in lines:
                content += f"{line.get('speaker')}: {line.get('text')}\n"
        except:
            pass
            
    response = HttpResponse(content, content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename="transcript_{meeting.id}.txt"'
    return response

from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse

@csrf_exempt
@login_required
def autosave_transcript(request, meeting_id):
    from core.models import Meeting, MeetingMinutes
    import json
    
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            transcript_text = data.get('transcript', '')
            mode = data.get('mode', 'replace')
            log_event = data.get('log_event', None)
            log_type = data.get('log_type', 'info')
            
            meeting = Meeting.objects.get(id=meeting_id)
            minutes, created = MeetingMinutes.objects.get_or_create(meeting=meeting, defaults={'notes': 'Autosaved transcript'})
            
            if log_event:
                import datetime
                from django.utils import timezone
                log_entry = {
                    'timestamp': timezone.now().strftime("%Y-%m-%dT%H:%M:%S"),
                    'event': log_event,
                    'type': log_type
                }
                current_logs = minutes.processing_logs if isinstance(minutes.processing_logs, list) else []
                current_logs.append(log_entry)
                minutes.processing_logs = current_logs
                minutes.save()
                return JsonResponse({'status': 'success'})
            
            if not transcript_text:
                return JsonResponse({'status': 'empty'})
            
            # Parse lines simply
            lines = transcript_text.strip().split('\n')
            parsed_lines = []
            for line in lines:
                if ':' in line:
                    speaker, text = line.split(':', 1)
                    parsed_lines.append({'speaker': speaker.strip(), 'text': text.strip()})
                else:
                    parsed_lines.append({'speaker': 'Unknown', 'text': line.strip()})
            
            if mode == 'append':
                existing_history = []
                if minutes.history:
                    try:
                        existing_history = json.loads(minutes.history)
                    except:
                        pass
                existing_history.extend(parsed_lines)
                minutes.history = json.dumps(existing_history)
            else:
                minutes.history = json.dumps(parsed_lines)
                
            minutes.save()
            
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})
    return JsonResponse({'status': 'invalid method'})



@login_required
def latest_transcript(request, meeting_id):
    from core.models import Meeting
    import json
    import datetime
    from django.utils import timezone
    from django.http import JsonResponse
    
    try:
        meeting = Meeting.objects.get(id=meeting_id)
        
        # Check if still ongoing
        start_dt = datetime.datetime.combine(meeting.date, meeting.time)
        if timezone.is_naive(start_dt):
            start_dt = timezone.make_aware(start_dt, timezone.get_default_timezone())
        try:
            duration = int(meeting.duration)
        except:
            duration = 60
        end_dt = start_dt + datetime.timedelta(minutes=duration)
        now = timezone.localtime(timezone.now())
        is_ongoing = start_dt <= now <= end_dt
        
        lines = []
        processing_logs = []
        minutes_status = ''
        if hasattr(meeting, 'minutes') and meeting.minutes:
            minutes_status = meeting.minutes.status
            if meeting.minutes.history:
                try:
                    lines = json.loads(meeting.minutes.history)
                except:
                    pass
            if meeting.minutes.processing_logs:
                processing_logs = meeting.minutes.processing_logs
            
        return JsonResponse({'status': 'success', 'lines': lines, 'logs': processing_logs, 'is_ongoing': is_ongoing, 'minutes_status': minutes_status})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})

@login_required
def unread_notification_count(request):
    from django.http import JsonResponse
    count = request.user.notifications.filter(is_read=False).count()
    return JsonResponse({'status': 'success', 'count': count})

@login_required
def fetch_notifications(request):
    from django.http import JsonResponse
    notifications = request.user.notifications.order_by('-created_at')[:20]
    data = []
    for n in notifications:
        data.append({
            'id': n.id,
            'message': n.message,
            'type': n.notification_type,
            'is_read': n.is_read,
            'created_at': n.created_at.strftime('%b %d, %I:%M %p'),
            'meeting_id': n.related_meeting.id if n.related_meeting else None
        })
    return JsonResponse({'status': 'success', 'notifications': data})

@login_required
def mark_notification_read(request, notification_id):
    from django.http import JsonResponse
    from core.models import InAppNotification
    try:
        notification = InAppNotification.objects.get(id=notification_id, user=request.user)
        notification.is_read = True
        notification.save()
        return JsonResponse({'status': 'success'})
    except InAppNotification.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Notification not found'})

@login_required
def clear_all_notifications(request):
    from django.http import JsonResponse
    from core.models import InAppNotification
    InAppNotification.objects.filter(user=request.user).delete()
    return JsonResponse({'status': 'success'})

@csrf_exempt
def zoom_webhook(request):
    import json
    import threading
    from core.models import Meeting, MeetingMinutes, Participant, CustomUser, ActionItem, InAppNotification
    from django.http import JsonResponse
    from django.core.files.base import ContentFile
    from django.core.mail import EmailMultiAlternatives
    from django.conf import settings
    import datetime
    from django.utils import timezone
    from core.zoom_api import ZoomAPI
    
    if request.method == 'POST':
        try:
            payload = json.loads(request.body)
            event = payload.get('event')
            
            # Simulated Webhook Payload contains a meeting_id
            zoom_meeting_id = payload.get('payload', {}).get('object', {}).get('id')
            
            if not zoom_meeting_id:
                return JsonResponse({'status': 'ignored', 'reason': 'no meeting id'})
                
            # Find the meeting by checking the description for the Zoom ID
            meeting = Meeting.objects.filter(description__icontains=zoom_meeting_id).first()
            
            if not meeting:
                return JsonResponse({'status': 'ignored', 'reason': 'meeting not found in db'})
                
            if event == 'recording.transcript_completed':
                minutes, _ = MeetingMinutes.objects.get_or_create(meeting=meeting)
                minutes.status = 'PROCESSING'
                
                log_entry = {
                    'timestamp': timezone.now().strftime("%Y-%m-%dT%H:%M:%S"),
                    'event': 'Zoom Webhook: Transcript Received. Starting Automated AI Pipeline...',
                    'type': 'success'
                }
                current_logs = minutes.processing_logs if isinstance(minutes.processing_logs, list) else []
                current_logs.append(log_entry)
                minutes.processing_logs = current_logs
                minutes.save()
                
                def background_pipeline(meeting_id):
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.pagesizes import letter
                    import io
                    import docx
                    
                    m = Meeting.objects.get(id=meeting_id)
                    mm = m.minutes
                    
                    try:
                        # 1. Fetch Mock Transcript
                        transcript_data = ZoomAPI.get_mock_transcript(m.title)
                        mm.history = json.dumps(transcript_data)
                        
                        log_entry = {'timestamp': timezone.now().strftime("%Y-%m-%dT%H:%M:%S"), 'event': 'Parsed Zoom VTT Transcript.', 'type': 'info'}
                        current_logs = mm.processing_logs if isinstance(mm.processing_logs, list) else []
                        current_logs.append(log_entry)
                        mm.processing_logs = current_logs
                        mm.save()
                        
                        # 2. AI Processing
                        mm.status = 'GENERATING'
                        mm.save()
                        
                        # Simulate AI Generation
                        mm.ai_summary = f"Automated Executive Summary for {m.title}. Discussed Q3 budget and agreed on 20% increase for social media."
                        mm.decisions = ["Approve 20% increase in social media budget."]
                        mm.action_items_raw = [{"task": "Update Excel spreadsheet", "owner": "Participant"}]
                        mm.status = 'READY'
                        
                        log_entry = {'timestamp': timezone.now().strftime("%Y-%m-%dT%H:%M:%S"), 'event': 'AI Summary & Actions Generated.', 'type': 'success'}
                        current_logs.append(log_entry)
                        mm.processing_logs = current_logs
                        mm.save()
                        
                        # Create Action Items in DB
                        admin_user = m.created_by
                        ActionItem.objects.create(
                            meeting=m,
                            task_name="Update Excel spreadsheet",
                            description="Automatically extracted from transcript.",
                            assigned_to=admin_user,
                            priority="HIGH",
                            due_date=timezone.now().date() + datetime.timedelta(days=1)
                        )
                        InAppNotification.objects.create(
                            user=admin_user,
                            message=f"You have been assigned an automated Action Item from {m.title}",
                            notification_type='TASK',
                            related_meeting=m
                        )
                        
                        # 3. Generate PDF
                        pdf_buffer = io.BytesIO()
                        p = canvas.Canvas(pdf_buffer, pagesize=letter)
                        p.drawString(100, 750, f"Automated MoM: {m.title}")
                        p.drawString(100, 730, f"Summary: {mm.ai_summary}")
                        p.showPage()
                        p.save()
                        
                        pdf_file = ContentFile(pdf_buffer.getvalue())
                        mm.pdf_export.save(f"automated_mom_{m.id}.pdf", pdf_file)
                        
                        # 4. Generate DOCX
                        doc_buffer = io.BytesIO()
                        doc = docx.Document()
                        doc.add_heading(f"Automated MoM: {m.title}", 0)
                        doc.add_paragraph(mm.ai_summary)
                        doc.save(doc_buffer)
                        
                        docx_file = ContentFile(doc_buffer.getvalue())
                        mm.docx_export.save(f"automated_mom_{m.id}.docx", docx_file)
                        
                        log_entry = {'timestamp': timezone.now().strftime("%Y-%m-%dT%H:%M:%S"), 'event': 'PDF & DOCX Generated successfully.', 'type': 'success'}
                        current_logs.append(log_entry)
                        mm.processing_logs = current_logs
                        mm.save()
                        
                        # 5. Email Participants
                        participants = [p.user for p in m.participants.all()]
                        if m.created_by not in participants:
                            participants.append(m.created_by)
                            
                        for user in participants:
                            subject = f"MoM Ready: {m.title}"
                            text_content = f"The automated Minutes of Meeting for '{m.title}' are now available.\n\nSummary:\n{mm.ai_summary}\n\nPlease check the portal to download the attachments."
                            msg = EmailMultiAlternatives(subject, text_content, settings.DEFAULT_FROM_EMAIL, [user.email])
                            
                            pdf_buffer.seek(0)
                            msg.attach(f"MoM_{m.id}.pdf", pdf_buffer.read(), 'application/pdf')
                            
                            try:
                                msg.send(fail_silently=True)
                            except Exception:
                                pass
                            
                            InAppNotification.objects.create(
                                user=user,
                                message=f"The Automated Minutes of Meeting for {m.title} are ready.",
                                notification_type='MOM',
                                related_meeting=m
                            )
                            
                        log_entry = {'timestamp': timezone.now().strftime("%Y-%m-%dT%H:%M:%S"), 'event': 'Email Notifications Sent.', 'type': 'success'}
                        current_logs.append(log_entry)
                        mm.processing_logs = current_logs
                        mm.save()
                        
                    except Exception as e:
                        log_entry = {'timestamp': timezone.now().strftime("%Y-%m-%dT%H:%M:%S"), 'event': f'Automated Pipeline Failed: {str(e)}', 'type': 'error'}
                        mm.processing_logs.append(log_entry)
                        mm.status = 'PENDING'
                        mm.save()
                        
                threading.Thread(target=background_pipeline, args=(meeting.id,)).start()
                return JsonResponse({'status': 'processing_started'})

            return JsonResponse({'status': 'ignored', 'event': event})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})
            
    return JsonResponse({'status': 'invalid method'})

@login_required
def simulate_zoom_webhook(request, meeting_id):
    import requests
    from core.models import Meeting
    from django.urls import reverse
    from django.contrib import messages
    from django.shortcuts import redirect
    
    meeting = Meeting.objects.get(id=meeting_id)
    
    # Extract zoom ID from description if present
    zoom_id = str(meeting.id)
    if meeting.description and "Zoom Meeting ID:" in meeting.description:
        import re
        match = re.search(r'Zoom Meeting ID:\s*(\d+)', meeting.description)
        if match:
            zoom_id = match.group(1)
            
    payload = {
        "event": "recording.transcript_completed",
        "payload": {
            "object": {
                "id": zoom_id,
                "topic": meeting.title
            }
        }
    }
    
    webhook_url = request.build_absolute_uri(reverse('zoom_webhook'))
    try:
        # Fire and forget request to self
        threading.Thread(target=requests.post, args=(webhook_url,), kwargs={'json': payload}).start()
        messages.success(request, "Zoom Webhook simulation triggered. Background pipeline is processing the transcript.")
    except Exception as e:
        messages.error(request, f"Failed to simulate webhook: {e}")
        
    return redirect('mom_detail', meeting_id=meeting.id)

@login_required
def upload_audio_for_transcription(request, meeting_id):
    from core.models import Meeting, MeetingMinutes
    from django.http import JsonResponse
    from core.ai_processor import start_ai_processing
    from django.utils import timezone
    import json
    
    if request.method == 'POST':
        try:
            meeting = Meeting.objects.get(id=meeting_id)
            audio_file = request.FILES.get('audio_file')
            
            if not audio_file:
                return JsonResponse({'status': 'error', 'message': 'No audio file provided'})
                
            minutes, _ = MeetingMinutes.objects.get_or_create(meeting=meeting)
            minutes.audio_file = audio_file
            minutes.status = 'PROCESSING'
            minutes.history = ''
            minutes.speaker_metrics = {}
            minutes.ai_summary = ''
            minutes.decisions = []
            minutes.action_items_raw = []
            minutes.processing_logs = [{'timestamp': timezone.now().isoformat(), 'message': 'Audio uploaded successfully. Enqueued for processing.'}]
            minutes.save()
            
            # Start background processing
            start_ai_processing(minutes.id)
            
            return JsonResponse({'status': 'success', 'message': 'Upload successful. Processing started.'})
            
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})
            
    return JsonResponse({'status': 'error', 'message': 'Invalid request'})

@login_required
def paste_transcript_text(request, meeting_id):
    from core.models import Meeting, MeetingMinutes
    from django.http import JsonResponse
    from core.ai_processor import start_ai_text_processing
    from django.utils import timezone
    
    if request.method == 'POST':
        try:
            meeting = Meeting.objects.get(id=meeting_id)
            transcript_text = request.POST.get('transcript_text')
            
            if not transcript_text or not transcript_text.strip():
                return JsonResponse({'status': 'error', 'message': 'No text provided'})
                
            minutes, _ = MeetingMinutes.objects.get_or_create(meeting=meeting)
            
            # Store the raw text in notes temporarily so the background worker can pick it up
            minutes.notes = transcript_text
            minutes.status = 'PROCESSING'
            minutes.history = ''
            minutes.speaker_metrics = {}
            minutes.ai_summary = ''
            minutes.decisions = []
            minutes.action_items_raw = []
            minutes.processing_logs = [{'timestamp': timezone.now().isoformat(), 'message': 'Text transcript received. Enqueued for processing.'}]
            minutes.save()
            
            # Start background processing
            start_ai_text_processing(minutes.id)
            
            return JsonResponse({'status': 'success', 'message': 'Upload successful. Processing started.'})
            
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})
            
    return JsonResponse({'status': 'error', 'message': 'Invalid request'})

@login_required
def check_transcription_status(request, meeting_id):
    from core.models import MeetingMinutes
    from django.http import JsonResponse
    try:
        minutes = MeetingMinutes.objects.get(meeting_id=meeting_id)
        return JsonResponse({
            'status': 'success',
            'state': minutes.status,
            'logs': minutes.processing_logs
        })
    except MeetingMinutes.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Not found'})

@login_required
def export_transcript(request, meeting_id, format):
    from core.models import MeetingMinutes
    from django.http import HttpResponse
    from django.shortcuts import get_object_or_404
    import json
    import io
    
    minutes = get_object_or_404(MeetingMinutes, meeting_id=meeting_id)
    transcript_data = json.loads(minutes.history) if minutes.history else []
    
    if format == 'txt':
        content = f"Transcript for {minutes.meeting.title}\n"
        content += f"Date: {minutes.meeting.date} | Time: {minutes.meeting.time}\n"
        content += "="*50 + "\n\n"
        
        for segment in transcript_data:
            content += f"{segment.get('timestamp', '')}\n{segment.get('speaker', 'Unknown')}:\n{segment.get('text', '')}\n\n"
            
        response = HttpResponse(content, content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="Transcript_{minutes.meeting.id}.txt"'
        return response
        
    elif format == 'docx':
        from docx import Document
        document = Document()
        document.add_heading(f"Transcript for {minutes.meeting.title}", 0)
        document.add_paragraph(f"Date: {minutes.meeting.date} | Time: {minutes.meeting.time}")
        
        for segment in transcript_data:
            p = document.add_paragraph()
            p.add_run(f"{segment.get('timestamp', '')}\n").bold = True
            p.add_run(f"{segment.get('speaker', 'Unknown')}:\n").bold = True
            p.add_run(segment.get('text', ''))
            
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Transcript_{minutes.meeting.id}.docx"'
        return response
        
    elif format == 'pdf':
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        c.setFont("Helvetica-Bold", 16)
        c.drawString(1*inch, height - 1*inch, f"Transcript for {minutes.meeting.title}")
        
        c.setFont("Helvetica", 12)
        c.drawString(1*inch, height - 1.3*inch, f"Date: {minutes.meeting.date} | Time: {minutes.meeting.time}")
        
        textobject = c.beginText(1*inch, height - 1.8*inch)
        textobject.setFont("Helvetica", 10)
        textobject.setLeading(14)
        
        for segment in transcript_data:
            if textobject.getY() < 1*inch:
                c.drawText(textobject)
                c.showPage()
                textobject = c.beginText(1*inch, height - 1*inch)
                textobject.setFont("Helvetica", 10)
                textobject.setLeading(14)
                
            textobject.setFont("Helvetica-Bold", 10)
            textobject.textLine(f"{segment.get('timestamp', '')} - {segment.get('speaker', 'Unknown')}:")
            textobject.setFont("Helvetica", 10)
            
            # Simple text wrap
            words = segment.get('text', '').split()
            line = ""
            for word in words:
                if len(line) + len(word) > 80:
                    textobject.textLine(line)
                    line = word + " "
                else:
                    line += word + " "
            textobject.textLine(line)
            textobject.textLine("")
            
        c.drawText(textobject)
        c.showPage()
        c.save()
        
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Transcript_{minutes.meeting.id}.pdf"'
        return response
        
    return JsonResponse({'status': 'error', 'message': 'Invalid format'})

@login_required
def retry_transcription_view(request, meeting_id):
    from core.models import Meeting, MeetingMinutes
    from django.shortcuts import redirect
    from django.contrib import messages
    from core.ai_processor import start_ai_processing, start_ai_text_processing
    import django.utils.timezone as timezone
    
    if request.method == 'POST':
        meeting = Meeting.objects.filter(id=meeting_id).first()
        if not meeting:
            messages.error(request, "Meeting not found.")
            return redirect('meetings')
            
        minutes, created = MeetingMinutes.objects.get_or_create(meeting=meeting)
        
        if not minutes.audio_file and not minutes.notes:
            messages.error(request, "No transcript found for this meeting. Please upload an audio file first.")
            return redirect('mom_detail', meeting_id=meeting_id)
            
        minutes.status = 'PENDING'
        minutes.processing_logs = [{'timestamp': timezone.now().isoformat(), 'message': 'Retrying AI processing...'}]
        minutes.save()
        
        if minutes.audio_file:
            start_ai_processing(minutes.id)
        elif minutes.notes:
            start_ai_text_processing(minutes.id)
            
    return redirect('mom_detail', meeting_id=meeting_id)

@login_required
def delete_transcription_view(request, meeting_id):
    from core.models import MeetingMinutes
    from django.shortcuts import get_object_or_404, redirect
    
    if request.method == 'POST':
        minutes = get_object_or_404(MeetingMinutes, meeting__id=meeting_id)
        # Clear data
        minutes.history = ""
        minutes.ai_summary = ""
        minutes.decisions = []
        minutes.action_items_raw = []
        minutes.processing_logs = []
        minutes.status = 'PENDING'
        if minutes.audio_file:
            minutes.audio_file.delete()
        minutes.save()
        
        # Delete related action items generated from MoM
        minutes.meeting.action_items.filter(description__icontains="Auto-generated action item").delete()
        
    return redirect('mom')
